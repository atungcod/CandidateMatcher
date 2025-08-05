import streamlit as st
from typing import Union
import PyPDF2
import docx
import io

class FileProcessor:
    """Handle processing of different file formats for resume content extraction"""
    
    def __init__(self):
        self.supported_formats = ['pdf', 'txt', 'docx']
    
    def process_file(self, uploaded_file) -> str:
        """
        Process uploaded file and extract text content
        
        Args:
            uploaded_file: Streamlit uploaded file object
            
        Returns:
            str: Extracted text content
            
        Raises:
            Exception: If file processing fails
        """
        try:
            file_extension = uploaded_file.name.split('.')[-1].lower()
            
            if file_extension not in self.supported_formats:
                raise ValueError(f"Unsupported file format: {file_extension}")
            
            if file_extension == 'pdf':
                return self._extract_pdf_text(uploaded_file)
            elif file_extension == 'txt':
                return self._extract_txt_text(uploaded_file)
            elif file_extension == 'docx':
                return self._extract_docx_text(uploaded_file)
            else:
                raise ValueError(f"Handler not implemented for: {file_extension}")
                
        except Exception as e:
            raise Exception(f"Failed to process file {uploaded_file.name}: {str(e)}")
    
    def _extract_pdf_text(self, uploaded_file) -> str:
        """Extract text from PDF file"""
        try:
            # Create a BytesIO object from the uploaded file
            pdf_bytes = io.BytesIO(uploaded_file.read())
            
            # Create PDF reader object
            pdf_reader = PyPDF2.PdfReader(pdf_bytes)
            
            # Extract text from all pages
            text_content = []
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text_content.append(page.extract_text())
            
            full_text = '\n'.join(text_content)
            
            if not full_text.strip():
                raise ValueError("No text content found in PDF")
            
            return full_text
            
        except Exception as e:
            raise Exception(f"PDF processing error: {str(e)}")
    
    def _extract_txt_text(self, uploaded_file) -> str:
        """Extract text from TXT file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    # Reset file pointer
                    uploaded_file.seek(0)
                    content = uploaded_file.read().decode(encoding)
                    
                    if content.strip():
                        return content
                        
                except UnicodeDecodeError:
                    continue
            
            raise ValueError("Could not decode text file with any supported encoding")
            
        except Exception as e:
            raise Exception(f"TXT processing error: {str(e)}")
    
    def _extract_docx_text(self, uploaded_file) -> str:
        """Extract text from DOCX file"""
        try:
            # Create a BytesIO object from the uploaded file
            docx_bytes = io.BytesIO(uploaded_file.read())
            
            # Load the document
            doc = docx.Document(docx_bytes)
            
            # Extract text from all paragraphs
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)
            
            full_text = '\n'.join(text_content)
            
            if not full_text.strip():
                raise ValueError("No text content found in DOCX")
            
            return full_text
            
        except Exception as e:
            raise Exception(f"DOCX processing error: {str(e)}")
    
    def validate_file_size(self, uploaded_file, max_size_mb: int = 10) -> bool:
        """
        Validate if file size is within acceptable limits
        
        Args:
            uploaded_file: Streamlit uploaded file object
            max_size_mb: Maximum file size in MB
            
        Returns:
            bool: True if file size is acceptable
        """
        file_size_mb = uploaded_file.size / (1024 * 1024)
        return file_size_mb <= max_size_mb
    
    def get_file_info(self, uploaded_file) -> dict:
        """
        Get information about uploaded file
        
        Args:
            uploaded_file: Streamlit uploaded file object
            
        Returns:
            dict: File information including name, size, type
        """
        return {
            'name': uploaded_file.name,
            'size_mb': round(uploaded_file.size / (1024 * 1024), 2),
            'type': uploaded_file.type,
            'extension': uploaded_file.name.split('.')[-1].lower()
        }
